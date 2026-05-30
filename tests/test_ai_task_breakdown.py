from datetime import datetime, timezone
from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.ai_task_breakdown import breakdown_requirements, breakdown_requirements_locally, normalize_breakdown_items
from app.database import init_db
from app.main import app
from app.repository import create_user, list_tasks
from app.schemas import TaskBreakdownItem
from app.settings import settings


client = TestClient(app)


def _hdr(user_id: int) -> dict:
    return {"X-User-Id": str(user_id)}


def _bootstrap() -> tuple[int, int, int]:
    init_db()
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S%f")
    admin = create_user("AI Admin", f"ai.admin.{stamp}@example.com", "admin", "PMO")
    manager = create_user("AI Manager", f"ai.manager.{stamp}@example.com", "manager", "PMO")
    staff = create_user("AI Staff", f"ai.staff.{stamp}@example.com", "staff", "Engineering")
    return int(admin["id"]), int(manager["id"]), int(staff["id"])


@pytest.fixture
def disable_ai(monkeypatch):
    monkeypatch.setattr(settings, "ai_api_key", "")


def test_local_breakdown_extracts_actionable_tasks() -> None:
    text = """
    He thong can cho phep quan ly tao task, keo tha Kanban va xuat bao cao KPI.
    Yeu cau tich hop Microsoft Teams de nguoi dung mo dashboard trong tab Teams.
    """

    items = breakdown_requirements_locally(text, max_tasks=5)

    assert len(items) >= 2
    assert all(item.title for item in items)
    assert {item.difficulty for item in items}.issubset({"easy", "medium", "hard"})
    assert all(1 <= item.story_points <= 13 for item in items)
    assert all(item.business_goal for item in items)
    assert all(item.acceptance_criteria for item in items)
    assert all(item.test_cases for item in items)


def test_local_breakdown_dashboard_kpi_uses_vietnamese_domain_details() -> None:
    text = "Tạo dashboard thống kê tiến độ task và KPI cho nhóm 3 người trong sprint hiện tại."

    items = breakdown_requirements_locally(text, max_tasks=3)

    assert len(items) == 1
    item = items[0]
    detail_text = " ".join(
        [
            item.business_goal,
            *item.subtasks,
            *item.acceptance_criteria,
            *item.data_requirements,
            *item.ui_components,
            *item.test_cases,
            item.demo_value,
        ]
    )
    assert item.type == "dashboard"
    assert "Tạo dashboard thống kê tiến độ task và KPI cho nhóm 3 người trong sprint hiện tại" in item.business_goal
    assert "..." not in item.business_goal
    assert "tổng task" in detail_text
    assert "task đúng hạn" in detail_text
    assert "task trễ hạn" in detail_text
    assert "điểm KPI từng thành viên" in detail_text
    assert "nhóm 3 người" in detail_text
    assert not any(english in detail_text for english in ("Clarify", "Confirm scope", "Implement", "Validate the happy path"))


def test_ai_preview_table_has_type_and_detail_headers() -> None:
    script = open("app/static/app.js", encoding="utf-8").read()

    assert "<th>Chọn</th><th>Task</th><th>Loại</th><th>SP</th><th>Deadline</th><th>Chi tiết</th>" in script
    assert "<th>Chọn</th><th>Task</th><th>Độ khó</th><th>SP</th><th>Deadline</th><th>Lý do</th>" not in script


def test_normalize_breakdown_items_preserves_structured_fields() -> None:
    items = normalize_breakdown_items(
        [
            {
                "title": "Build KPI drilldown",
                "type": "dashboard",
                "description": "Show overdue task details from KPI cards.",
                "business_goal": "Help managers explain KPI score changes.",
                "subtasks": ["Define drilldown query", "Render task table"],
                "acceptance_criteria": ["Manager can open overdue task list"],
                "data_requirements": ["tasks.deadline", "tasks.status"],
                "ui_components": ["KPI card", "Drawer table"],
                "test_cases": ["No overdue tasks shows empty state"],
                "dependencies": ["KPI endpoint"],
                "risks": ["Large task lists may be slow"],
                "demo_value": "Shows auditable KPI evidence.",
                "suggested_role": "Manager",
                "story_points": 5,
                "difficulty": "medium",
                "deadline_offset_days": 7,
                "rationale": "Important for demo.",
            }
        ],
        max_tasks=3,
    )

    assert len(items) == 1
    item = items[0]
    assert item.type == "dashboard"
    assert item.business_goal == "Help managers explain KPI score changes."
    assert item.subtasks == ["Define drilldown query", "Render task table"]
    assert item.acceptance_criteria == ["Manager can open overdue task list"]
    assert item.data_requirements == ["tasks.deadline", "tasks.status"]
    assert item.ui_components == ["KPI card", "Drawer table"]
    assert item.test_cases == ["No overdue tasks shows empty state"]
    assert item.dependencies == ["KPI endpoint"]
    assert item.risks == ["Large task lists may be slow"]
    assert item.demo_value == "Shows auditable KPI evidence."
    assert item.suggested_role == "Manager"


def test_ai_breakdown_uses_fallback_model_when_primary_fails(monkeypatch) -> None:
    monkeypatch.setattr(settings, "ai_provider", "openai_compatible")
    monkeypatch.setattr(settings, "ai_api_key", "ollama")
    monkeypatch.setattr(settings, "ai_model", "qwen3:8b")
    monkeypatch.setattr(settings, "ai_fallback_model", "qwen2.5:7b")
    called_models: list[str | None] = []

    def fake_breakdown(text: str, project_context: str | None, max_tasks: int, model: str | None = None):
        called_models.append(model)
        if model == "qwen3:8b":
            raise RuntimeError("primary unavailable")
        return [
            TaskBreakdownItem(
                title="Fallback task",
                description=text,
                story_points=3,
                difficulty="medium",
                deadline_offset_days=7,
            )
        ]

    monkeypatch.setattr("app.ai_task_breakdown._breakdown_with_openai_compatible", fake_breakdown)

    result = breakdown_requirements("Tao bao cao KPI", max_tasks=3)

    assert result.source == "openai_compatible"
    assert result.items[0].title == "Fallback task"
    assert called_models == ["qwen3:8b", "qwen2.5:7b"]
    assert any("trying fallback model qwen2.5:7b" in warning for warning in result.warnings)


def test_task_breakdown_preview_review_and_import_flow(disable_ai) -> None:
    _admin_id, manager_id, staff_id = _bootstrap()
    before_count = len(list_tasks(assignee_id=staff_id))
    text = """
    Xay dung module AI phan ra yeu cau tu tai lieu docx.
    Nguoi quan ly co the xem danh sach task de xuat va chon task de import vao Kanban.
    He thong can ghi audit log sau khi tao task.
    """

    preview = client.post(
        "/ai/task-breakdown",
        headers=_hdr(manager_id),
        json={"text": text, "max_tasks": 4},
    )
    assert preview.status_code == 200
    payload = preview.json()
    assert payload["source"] in {"heuristic", "openai_compatible"}
    assert payload["ai_draft_id"]
    assert payload["status"] == "draft"
    assert len(payload["items"]) >= 1
    assert payload["items"][0]["title"]
    assert len(list_tasks(assignee_id=staff_id)) == before_count

    drafts = client.get("/ai/task-breakdown/drafts", headers=_hdr(manager_id))
    assert drafts.status_code == 200
    assert any(int(item["id"]) == int(payload["ai_draft_id"]) for item in drafts.json())

    edited_items = payload["items"][:2]
    edited_items[0]["title"] = "Reviewed AI task"
    edited_items[0]["type"] = "integration"
    edited_items[0]["business_goal"] = "Manager can approve AI generated work packages."
    edited_items[0]["subtasks"] = ["Validate draft items", "Import selected tasks"]
    edited_items[0]["acceptance_criteria"] = ["Selected draft items create Kanban tasks"]
    edited_items[0]["data_requirements"] = ["ai_task_drafts.generated_tasks"]
    edited_items[0]["ui_components"] = ["AI draft review drawer"]
    edited_items[0]["test_cases"] = ["Import skips unselected items"]
    edited_items[0]["dependencies"] = ["RBAC ai.import permission"]
    edited_items[0]["risks"] = ["Reviewer may import incomplete scope"]
    edited_items[0]["demo_value"] = "Shows human approval before task creation."
    edited_items[0]["suggested_role"] = "Manager"
    review_resp = client.patch(
        f"/ai/task-breakdown/drafts/{payload['ai_draft_id']}/review",
        headers=_hdr(manager_id),
        json={
            "items": edited_items,
            "review_note": "Manager approved the useful tasks.",
            "edit_reason": "Trimmed scope and renamed first task.",
        },
    )
    assert review_resp.status_code == 200
    reviewed = review_resp.json()
    assert reviewed["status"] == "reviewed"
    assert reviewed["reviewer_id"] == manager_id
    assert reviewed["review_note"] == "Manager approved the useful tasks."
    assert reviewed["edit_reason"] == "Trimmed scope and renamed first task."
    assert reviewed["items"][0]["title"] == "Reviewed AI task"
    assert reviewed["items"][0]["subtasks"] == ["Validate draft items", "Import selected tasks"]

    import_resp = client.post(
        "/ai/task-breakdown/import",
        headers=_hdr(manager_id),
        json={
            "ai_draft_id": payload["ai_draft_id"],
            "assignee_id": staff_id,
            "project_id": None,
            "sprint_id": None,
        },
    )
    assert import_resp.status_code == 200
    imported = import_resp.json()
    assert imported["created_count"] == len(edited_items)
    assert len(imported["tasks"]) == len(edited_items)
    assert all(task["assignee_id"] == staff_id for task in imported["tasks"])
    assert imported["tasks"][0]["title"] == "Reviewed AI task"

    task_detail = client.get(f"/tasks/{imported['tasks'][0]['id']}", headers=_hdr(manager_id))
    assert task_detail.status_code == 200
    ai_detail = task_detail.json()["ai_detail"]
    assert ai_detail["source_ai_draft_id"] == payload["ai_draft_id"]
    assert ai_detail["type"] == "integration"
    assert ai_detail["business_goal"] == "Manager can approve AI generated work packages."
    assert ai_detail["subtasks"] == ["Validate draft items", "Import selected tasks"]
    assert ai_detail["acceptance_criteria"] == ["Selected draft items create Kanban tasks"]
    assert ai_detail["test_cases"] == ["Import skips unselected items"]

    imported_draft = client.get(
        f"/ai/task-breakdown/drafts/{payload['ai_draft_id']}",
        headers=_hdr(manager_id),
    )
    assert imported_draft.status_code == 200
    assert imported_draft.json()["status"] == "imported"
    assert imported_draft.json()["imported_at"]

    import_again = client.post(
        "/ai/task-breakdown/import",
        headers=_hdr(manager_id),
        json={"ai_draft_id": payload["ai_draft_id"], "assignee_id": staff_id},
    )
    assert import_again.status_code == 400


def test_task_breakdown_docx_creates_draft(disable_ai) -> None:
    _admin_id, manager_id, _staff_id = _bootstrap()
    doc = Document()
    doc.add_paragraph("Xay dung module import AI draft cho manager review truoc khi tao task that.")
    doc.add_paragraph("Can co audit log, trang thai reviewed va chan import lai batch da imported.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    preview = client.post(
        "/ai/task-breakdown/docx",
        headers=_hdr(manager_id),
        files={
            "file": (
                "ai-draft-review.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"max_tasks": "3"},
    )

    assert preview.status_code == 200
    payload = preview.json()
    assert payload["ai_draft_id"]
    assert payload["status"] == "draft"
    detail = client.get(f"/ai/task-breakdown/drafts/{payload['ai_draft_id']}", headers=_hdr(manager_id))
    assert detail.status_code == 200
    assert detail.json()["source_type"] == "docx"
    assert detail.json()["source_name"] == "ai-draft-review.docx"


def test_task_breakdown_docx_does_not_query_rag_when_disabled(monkeypatch, disable_ai) -> None:
    _admin_id, manager_id, _staff_id = _bootstrap()
    doc = Document()
    doc.add_paragraph("Xay dung docx endpoint khong goi RAG khi use_rag bi tat.")
    doc.add_paragraph("Noi dung du dai de tao task breakdown bang heuristic fallback.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    def fail_query_rag(*_args, **_kwargs):
        raise AssertionError("query_rag should not be called")

    monkeypatch.setattr("app.routers.ai.query_rag", fail_query_rag)

    preview = client.post(
        "/ai/task-breakdown/docx",
        headers=_hdr(manager_id),
        files={
            "file": (
                "no-rag.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"max_tasks": "3", "use_rag": "false"},
    )

    assert preview.status_code == 200
    assert preview.json()["retrieved_context_count"] == 0


def test_task_breakdown_docx_queries_rag_with_custom_query_when_enabled(monkeypatch, disable_ai) -> None:
    _admin_id, manager_id, _staff_id = _bootstrap()
    doc = Document()
    doc.add_paragraph("Xay dung docx endpoint goi RAG khi use_rag duoc bat.")
    doc.add_paragraph("Noi dung du dai de tao task breakdown bang heuristic fallback.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    called_queries: list[str] = []

    def fake_query_rag(query: str, limit: int = 5, current_user: dict | None = None) -> list[dict]:
        called_queries.append(query)
        assert limit == 5
        assert current_user is not None
        return [
            {
                "source_label": "custom-rag-source",
                "content": "Context tu RAG cho docx task breakdown.",
                "score": 1.0,
            }
        ]

    monkeypatch.setattr("app.routers.ai.query_rag", fake_query_rag)

    preview = client.post(
        "/ai/task-breakdown/docx",
        headers=_hdr(manager_id),
        files={
            "file": (
                "with-rag.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"max_tasks": "3", "use_rag": "true", "rag_query": "custom query from form"},
    )

    assert preview.status_code == 200
    assert called_queries == ["custom query from form"]
    assert preview.json()["retrieved_context_count"] == 1
    assert "custom-rag-source" in preview.json()["retrieved_sources"]


def test_staff_cannot_generate_or_import_ai_tasks(disable_ai) -> None:
    _admin_id, _manager_id, staff_id = _bootstrap()

    preview = client.post(
        "/ai/task-breakdown",
        headers=_hdr(staff_id),
        json={"text": "Tao dashboard KPI", "max_tasks": 3},
    )
    assert preview.status_code == 403

    drafts = client.get("/ai/task-breakdown/drafts", headers=_hdr(staff_id))
    assert drafts.status_code == 403

    review = client.patch(
        "/ai/task-breakdown/drafts/1/review",
        headers=_hdr(staff_id),
        json={"items": [{"title": "Nope", "difficulty": "medium", "story_points": 3, "deadline_offset_days": 7}]},
    )
    assert review.status_code == 403

    import_resp = client.post(
        "/ai/task-breakdown/import",
        headers=_hdr(staff_id),
        json={"ai_draft_id": 1, "assignee_id": staff_id},
    )
    assert import_resp.status_code == 403
