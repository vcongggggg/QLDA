from __future__ import annotations

import io
import json
import re
import time
from dataclasses import dataclass

import httpx
from docx import Document

from app.schemas import TaskBreakdownItem
from app.settings import settings


MAX_DOCX_BYTES = 2 * 1024 * 1024
ALLOWED_STORY_POINTS = (1, 2, 3, 5, 8, 13)


@dataclass(frozen=True)
class BreakdownResult:
    source: str
    items: list[TaskBreakdownItem]
    warnings: list[str]


def extract_docx_text(data: bytes) -> str:
    if len(data) > MAX_DOCX_BYTES:
        raise ValueError("docx file is too large; max size is 2MB")
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def normalize_breakdown_items(raw_items: list[dict], max_tasks: int) -> list[TaskBreakdownItem]:
    output: list[TaskBreakdownItem] = []
    for index, raw in enumerate(raw_items[:max_tasks], start=1):
        title = str(raw.get("title") or raw.get("name") or "").strip()
        if not title:
            continue
        description = _clean_text(raw.get("description"), max_length=1200)
        difficulty = str(raw.get("difficulty") or "medium").strip().lower()
        if difficulty not in {"easy", "medium", "hard"}:
            difficulty = "medium"
        try:
            story_points = int(raw.get("story_points", 3))
        except (TypeError, ValueError):
            story_points = 3
        if story_points not in ALLOWED_STORY_POINTS:
            story_points = min(ALLOWED_STORY_POINTS, key=lambda value: abs(value - story_points))
        try:
            offset = int(raw.get("deadline_offset_days", 7 + index))
        except (TypeError, ValueError):
            offset = 7 + index
        defaults = _structured_defaults(raw, title=title, description=description)
        output.append(
            TaskBreakdownItem(
                title=title[:200],
                type=_clean_text(raw.get("type"), max_length=80) or "implementation",
                description=description,
                business_goal=_clean_text(raw.get("business_goal"), max_length=1200) or defaults["business_goal"],
                subtasks=_string_list(raw.get("subtasks")) or defaults["subtasks"],
                acceptance_criteria=_string_list(raw.get("acceptance_criteria")) or defaults["acceptance_criteria"],
                data_requirements=_string_list(raw.get("data_requirements")) or defaults["data_requirements"],
                ui_components=_string_list(raw.get("ui_components")) or defaults["ui_components"],
                test_cases=_string_list(raw.get("test_cases")) or defaults["test_cases"],
                dependencies=_string_list(raw.get("dependencies")) or defaults["dependencies"],
                risks=_string_list(raw.get("risks")) or defaults["risks"],
                demo_value=_clean_text(raw.get("demo_value"), max_length=1200) or defaults["demo_value"],
                suggested_role=_clean_text(raw.get("suggested_role"), max_length=120) or "Manager",
                story_points=story_points,
                difficulty=difficulty,
                deadline_offset_days=max(1, min(offset, 90)),
                rationale=_clean_text(raw.get("rationale"), max_length=500),
                selected=bool(raw.get("selected", True)),
            )
        )
    return output


def breakdown_requirements_locally(text: str, max_tasks: int = 8) -> list[TaskBreakdownItem]:
    candidates = _candidate_sentences(text)
    if not candidates:
        candidates = [text.strip()]
    raw_items: list[dict] = []
    for index, sentence in enumerate(candidates[:max_tasks], start=1):
        title = _title_from_sentence(sentence)
        local_defaults = _local_fallback_fields(sentence, title)
        raw_items.append(
            {
                "title": title,
                "description": sentence.strip(),
                "story_points": _estimate_story_points(sentence),
                "difficulty": _estimate_difficulty(sentence),
                "deadline_offset_days": min(90, 5 + index * 2),
                "type": _infer_task_type(sentence),
                **local_defaults,
            }
        )
    return normalize_breakdown_items(raw_items, max_tasks=max_tasks)


def breakdown_requirements(text: str, project_context: str | None = None, max_tasks: int = 8) -> BreakdownResult:
    warnings: list[str] = []
    if settings.ai_provider == "openai_compatible" and settings.ai_api_key:
        last_exc: Exception | None = None
        models = [settings.ai_model]
        if settings.ai_fallback_model and settings.ai_fallback_model not in models:
            models.append(settings.ai_fallback_model)

        for model_index, model in enumerate(models):
            for attempt in range(1, 4):
                try:
                    items = _breakdown_with_openai_compatible(text, project_context, max_tasks, model=model)
                    if items:
                        return BreakdownResult(source="openai_compatible", items=items, warnings=warnings)
                    warnings.append(f"AI model {model} returned no tasks.")
                    break
                except httpx.HTTPStatusError as exc:
                    last_exc = exc
                    if exc.response.status_code == 429 and attempt < 3:
                        warnings.append(f"Rate limit on {model} (attempt {attempt}/3), retrying in {attempt * 3}s...")
                        time.sleep(attempt * 3)
                        continue
                    break
                except Exception as exc:  # noqa: BLE001 - fallback must be robust for demos
                    last_exc = exc
                    break
            if model_index < len(models) - 1:
                warnings.append(f"AI model {model} failed; trying fallback model {models[-1]}.")
                continue
        if last_exc is not None:
            warnings.append("AI provider failed; used local fallback.")
    else:
        warnings.append("AI_API_KEY is not configured; used local fallback.")
    return BreakdownResult(source="heuristic", items=breakdown_requirements_locally(text, max_tasks), warnings=warnings)


def _breakdown_with_openai_compatible(
    text: str,
    project_context: str | None,
    max_tasks: int,
    model: str | None = None,
) -> list[TaskBreakdownItem]:
    base_url = settings.ai_base_url.rstrip("/")
    prompt = {
        "role": "system",
        "content": (
            "You are an AI Project Manager Assistant inside TeamsWork, an internal task, sprint, KPI, "
            "reporting, Microsoft Teams, and AI task breakdown system. Your job is to turn requirements "
            "into work packages that can be assigned, reviewed, tested, and demonstrated immediately. "
            "Always answer in the same language as the input requirement. Return only valid JSON with key "
            "'tasks'. Each task must include: title, type, description, business_goal, subtasks, "
            "acceptance_criteria, data_requirements, ui_components, test_cases, dependencies, risks, "
            "demo_value, suggested_role, story_points, difficulty, deadline_offset_days, rationale. "
            "story_points must be one of 1,2,3,5,8,13. difficulty must be easy|medium|hard. "
            "If the requirement is about dashboard, report, or statistics, every task must mention metrics, "
            "data source, filters, charts or tables, empty state, and validation test. Do not create generic "
            "tasks like 'Design and Develop Dashboard'. Do not invent scope far from the user's requirement."
        ),
    }
    user_content = {
        "project_context": project_context or "TeamsWork task/KPI management system integrated with Microsoft Teams",
        "max_tasks": max_tasks,
        "requirements": text[:50000],
    }
    with httpx.Client(timeout=settings.ai_task_breakdown_timeout_seconds) as client:
        response = client.post(
            f"{base_url}/chat/completions",
            headers={
                "Authorization": f"Bearer {settings.ai_api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": model or settings.ai_model,
                "messages": [prompt, {"role": "user", "content": json.dumps(user_content, ensure_ascii=False)}],
                "temperature": 0.2,
                "response_format": {"type": "json_object"},
            },
        )
        response.raise_for_status()
    content = response.json()["choices"][0]["message"]["content"]
    payload = json.loads(content)
    raw_tasks = payload.get("tasks", payload if isinstance(payload, list) else [])
    if not isinstance(raw_tasks, list):
        return []
    return normalize_breakdown_items(raw_tasks, max_tasks=max_tasks)


def _candidate_sentences(text: str) -> list[str]:
    lines = [line.strip(" -•\t") for line in text.splitlines() if line.strip()]
    chunks: list[str] = []
    for line in lines:
        if len(line) <= 180:
            chunks.append(line)
        else:
            chunks.extend(part.strip() for part in re.split(r"(?<=[.!?。])\s+", line) if part.strip())
    keywords = ("cần", "cho phép", "quản lý", "tạo", "xem", "cập nhật", "xuất", "tích hợp", "đăng nhập", "phân quyền")
    prioritized = [chunk for chunk in chunks if any(keyword in chunk.lower() for keyword in keywords)]
    return prioritized or chunks


def _title_from_sentence(sentence: str) -> str:
    cleaned = re.sub(r"\s+", " ", sentence).strip(" .,:;")
    cleaned = re.sub(r"^(hệ thống|người dùng|admin|manager|quản lý)\s+", "", cleaned, flags=re.IGNORECASE)
    if len(cleaned) > 72:
        cleaned = cleaned[:69].rstrip() + "..."
    if not cleaned:
        return "Phân tích và triển khai yêu cầu"
    return cleaned[0].upper() + cleaned[1:]


def _estimate_difficulty(sentence: str) -> str:
    lower = sentence.lower()
    hard_markers = ("tích hợp", "sso", "azure", "teams", "báo cáo", "ai", "phân rã", "đồng bộ")
    easy_markers = ("hiển thị", "xem", "lọc", "danh sách")
    if any(marker in lower for marker in hard_markers):
        return "hard"
    if any(marker in lower for marker in easy_markers):
        return "easy"
    return "medium"


def _estimate_story_points(sentence: str) -> int:
    difficulty = _estimate_difficulty(sentence)
    if difficulty == "hard":
        return 8
    if difficulty == "easy":
        return 2
    return 5


def _clean_text(value: object, max_length: int) -> str | None:
    text = str(value or "").strip()
    if not text:
        return None
    return text[:max_length]


def _string_list(value: object, max_items: int = 12, max_length: int = 240) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        raw_items = re.split(r"[\n;]+", value)
    elif isinstance(value, list):
        raw_items = value
    else:
        raw_items = [value]
    items: list[str] = []
    for raw in raw_items:
        text = str(raw or "").strip(" -\t")
        if text:
            items.append(text[:max_length])
        if len(items) >= max_items:
            break
    return items


def _structured_defaults(raw: dict, title: str, description: str | None) -> dict[str, object]:
    base = description or title
    if _is_dashboard_request(base):
        return _dashboard_fallback_fields(base, title, vietnamese=_looks_vietnamese(base))
    if _looks_vietnamese(base):
        return {
            "business_goal": f"Chuyển yêu cầu thành kết quả có thể giao việc, review và demo: {base}",
            "subtasks": [f"Làm rõ phạm vi cho {title}", f"Triển khai luồng chính cho {title}", f"Chuẩn bị bằng chứng kiểm thử cho {title}"],
            "acceptance_criteria": [f"{title} có mô tả rõ, có thể review và bàn giao cho manager."],
            "data_requirements": ["Nội dung yêu cầu", "Bối cảnh dự án"],
            "ui_components": ["Drawer chi tiết task"],
            "test_cases": [f"Kiểm thử {title} với dữ liệu mẫu đại diện."],
            "dependencies": _string_list(raw.get("dependencies")) or ["Manager review"],
            "risks": _string_list(raw.get("risks")) or ["Yêu cầu có thể cần manager làm rõ thêm."],
            "demo_value": f"Cho thấy tiến độ xử lý {title}.",
        }
    return {
        "business_goal": f"Deliver a clear business outcome for: {base}",
        "subtasks": [f"Clarify scope for {title}", f"Implement {title}", f"Review and demo {title}"],
        "acceptance_criteria": [f"{title} is complete and reviewable by the manager."],
        "data_requirements": ["Requirement text", "Project context"],
        "ui_components": ["Task detail drawer"],
        "test_cases": [f"Validate {title} with representative input."],
        "dependencies": _string_list(raw.get("dependencies")),
        "risks": _string_list(raw.get("risks")) or ["Scope may need manager clarification."],
        "demo_value": f"Demonstrates progress on {title}.",
    }


def _infer_task_type(sentence: str) -> str:
    lower = sentence.lower()
    if _is_dashboard_request(sentence):
        return "dashboard"
    if any(marker in lower for marker in ("api", "integration", "teams", "sync", "tÃ­ch há»£p")):
        return "integration"
    if any(marker in lower for marker in ("test", "qa", "kiá»ƒm thá»­")):
        return "qa"
    return "implementation"


def _local_fallback_fields(sentence: str, title: str) -> dict[str, object]:
    if _is_dashboard_request(sentence):
        return _dashboard_fallback_fields(sentence, title, vietnamese=_looks_vietnamese(sentence))
    if _looks_vietnamese(sentence):
        return {
            "business_goal": f"Chuyển yêu cầu thành task rõ phạm vi, có thể giao việc và kiểm thử: {sentence.strip()}",
            "subtasks": [
                f"Làm rõ phạm vi và kết quả cần bàn giao cho {title}",
                f"Triển khai luồng nghiệp vụ chính cho {title}",
                f"Chuẩn bị bằng chứng kiểm thử và nội dung demo cho {title}",
            ],
            "acceptance_criteria": [
                "Manager đọc task và hiểu được kết quả cần bàn giao.",
                "Task có thể được review trước khi chuyển sang hoàn thành.",
            ],
            "data_requirements": ["Nội dung yêu cầu", "Bối cảnh dự án", "Người phụ trách"],
            "ui_components": ["Thẻ task Kanban", "Drawer chi tiết task"],
            "test_cases": [
                "Kiểm thử luồng chính với dữ liệu mẫu.",
                "Kiểm thử khi thiếu dữ liệu tùy chọn không làm hỏng luồng xử lý.",
            ],
            "dependencies": ["Manager review", "Người phụ trách khả dụng"],
            "risks": ["Yêu cầu có thể cần làm rõ thêm trước khi triển khai."],
            "demo_value": "Cho thấy AI chuyển yêu cầu tiếng Việt thành task có thể review và import.",
            "suggested_role": "Manager",
            "rationale": "Sinh tự động từ yêu cầu/tài liệu bằng heuristic fallback.",
        }
    return {
        "business_goal": f"Clarify and deliver the full requirement: {sentence.strip()}",
        "subtasks": [
            f"Confirm scope for {title}",
            f"Implement the core workflow for {title}",
            f"Prepare validation evidence for {title}",
        ],
        "acceptance_criteria": [
            "The assigned user can understand the expected outcome from the task description.",
            "The task can be reviewed by a manager before completion.",
        ],
        "data_requirements": ["Requirement text", "Project context", "Assigned user"],
        "ui_components": ["Kanban task card", "Task detail drawer"],
        "test_cases": [
            "Validate the happy path with representative data.",
            "Validate empty or missing optional data does not break the workflow.",
        ],
        "dependencies": ["Manager review", "Available assignee"],
        "risks": ["Requirement may need clarification before implementation."],
        "demo_value": "Shows AI converting requirements into reviewable implementation work.",
        "suggested_role": "Manager",
        "rationale": "Generated from requirement/document text by heuristic fallback.",
    }


def _dashboard_fallback_fields(sentence: str, title: str, vietnamese: bool) -> dict[str, object]:
    if vietnamese:
        return {
            "business_goal": f"Cung cấp dashboard giúp manager theo dõi tiến độ task và KPI từ yêu cầu: {sentence.strip()}",
            "subtasks": [
                "Xác định KPI cần hiển thị: tổng task, task đúng hạn, task trễ hạn, điểm KPI từng thành viên",
                "Thiết kế card tổng quan và biểu đồ tiến độ task theo trạng thái",
                "Xây dựng danh sách task trễ hạn kèm người phụ trách và deadline",
                "Tính hiệu suất từng thành viên trong nhóm 3 người",
                "Kiểm thử số liệu dashboard với dữ liệu mẫu của sprint hiện tại",
            ],
            "acceptance_criteria": [
                "Dashboard hiển thị tổng task, task đúng hạn, task trễ hạn và điểm KPI từng thành viên.",
                "Manager xem được tiến độ sprint hiện tại và drilldown danh sách task trễ hạn.",
                "Số liệu của nhóm 3 người khớp với dữ liệu task/KPI mẫu dùng trong kiểm thử.",
            ],
            "data_requirements": [
                "tasks.status",
                "tasks.deadline",
                "tasks.completed_at",
                "tasks.assignee_id",
                "story_points và difficulty để đối chiếu KPI",
                "dữ liệu KPI từng thành viên trong sprint hiện tại",
            ],
            "ui_components": [
                "Card tổng quan tổng task, đúng hạn, trễ hạn",
                "Biểu đồ tiến độ task theo trạng thái",
                "Bảng task trễ hạn",
                "Bảng KPI từng thành viên",
            ],
            "test_cases": [
                "Kiểm thử dashboard với nhóm 3 người có task đúng hạn, task trễ hạn và task chưa hoàn thành.",
                "Kiểm thử điểm KPI từng thành viên khớp công thức KPI hiện tại.",
                "Kiểm thử trạng thái rỗng khi sprint chưa có task.",
            ],
            "dependencies": ["Dữ liệu task của sprint hiện tại", "API KPI", "Quyền xem dashboard/KPI của manager"],
            "risks": ["Số liệu dashboard có thể lệch nếu filter sprint, deadline hoặc assignee không thống nhất với KPI."],
            "demo_value": "Demo rõ tiến độ task, task trễ hạn và KPI từng thành viên trong cùng một dashboard.",
            "suggested_role": "Manager",
            "rationale": "Sinh tự động từ yêu cầu dashboard/KPI bằng heuristic fallback theo tiếng Việt.",
        }
    return {
        "business_goal": f"Provide a dashboard for managers to track task progress and KPI from the full requirement: {sentence.strip()}",
        "subtasks": [
            "Define metrics: total tasks, on-time tasks, overdue tasks, and KPI score per member",
            "Design summary cards and task progress charts",
            "Build the overdue task list with assignee and deadline",
            "Calculate member performance for the requested team",
            "Validate dashboard numbers with representative sprint data",
        ],
        "acceptance_criteria": [
            "Dashboard shows task progress, overdue tasks, and KPI score per member.",
            "Manager can validate the dashboard numbers against task and KPI data.",
        ],
        "data_requirements": ["tasks.status", "tasks.deadline", "tasks.completed_at", "tasks.assignee_id", "KPI score data"],
        "ui_components": ["Summary cards", "Progress chart", "Overdue task table", "Member KPI table"],
        "test_cases": ["Validate dashboard metrics with sample sprint data.", "Validate empty sprint state."],
        "dependencies": ["Task data", "KPI API", "Manager dashboard permission"],
        "risks": ["Dashboard metrics can drift if filters differ from KPI calculation."],
        "demo_value": "Shows task progress and KPI evidence in one manager dashboard.",
        "suggested_role": "Manager",
        "rationale": "Generated from dashboard/KPI requirement by heuristic fallback.",
    }


def _is_dashboard_request(text: str) -> bool:
    lower = text.lower()
    return any(
        marker in lower
        for marker in (
            "dashboard",
            "report",
            "statistics",
            "progress",
            "kpi",
            "thống kê",
            "thong ke",
            "báo cáo",
            "bao cao",
            "tiến độ",
            "tien do",
        )
    )


def _looks_vietnamese(text: str) -> bool:
    lower = text.lower()
    if re.search(r"[ăâđêôơưáàảãạấầẩẫậắằẳẵặéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ]", lower):
        return True
    return any(
        marker in lower
        for marker in (
            "tao ",
            "xay dung",
            "yeu cau",
            "nguoi dung",
            "quan ly",
            "bao cao",
            "thong ke",
            "tien do",
            "nhom",
            "cong viec",
        )
    )
